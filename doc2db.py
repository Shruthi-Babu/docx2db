from docx import Document
from boltons import iterutils
import sqlite3
import re


conn = sqlite3.connect('Docs/5a.db')
conn2 = sqlite3.connect('Docs/2017e.db')
c=conn.cursor()
c2 = conn2.cursor()
c.execute('''CREATE TABLE IF NOT EXISTS classes (day TEXT, slot TEXT, room TEXT, mapid TEXT)''')
c.execute('''CREATE TABLE IF NOT EXISTS faculty(fid TEXT, name TEXT, tag TEXT,designation TEXT, qualification TEXT, phoneno TEXT, emailid TEXT)''')
c.execute('''CREATE TABLE IF NOT EXISTS slot(slotnumber TEXT,timings TEXT)''')
c.execute('''CREATE TABLE IF NOT EXISTS subjects(mapid TEXT, fid TEXT, sem TEXT,sub TEXT)''')

#c.execute('''SELECT * INTO 5a.slot FROM 2017e.slot''')

teacherArray=[]
mapid_prefix = "5a"
#slotArray = []
#timeArray = []
fidArray = []
tagArray = []
c2.execute('''SELECT timings FROM slot ''')
timeArray = c2.fetchall()

c2.execute('''SELECT slotnumber FROM slot ''')
slotArray = c2.fetchall()

#slotsdict = {}
facDict = {}

c2.execute('''SELECT fid FROM faculty ''')
fidArray = c2.fetchall()

c2.execute('''SELECT tag FROM faculty''')
tagArray = c2.fetchall()

# for i in range(0,13):
#     slotsdict[timeArray[i]] = slotArray[i]

for i in range(0,13):
    facDict[tagArray[i]] = fidArray[i]

#print(slotsdict)
# print facDict
# print len(slotsdict)
# print len(facDict)

wordDoc = Document("Docs/5a.docx")
text = []
time_tables = []
class_rooms = []
subjects = []
TABLE = [[]]
for info in wordDoc.paragraphs:
        text.append(info.text)
#print text
subArray=[]
subcodeArray=[]
for item in text:
    if "CS" in item and "." in item:
        arr = re.split("[\t]+", str(item))
        sub_code = re.split("[0-9]?\.", arr[0])[1]
        sub_name = arr[2].strip()
        sublist=sub_name.splitlines()
        subcodeList=sub_code.splitlines()
        #print sub_name
        #print sublist
        #print sub_code
        for subs in sublist:
            subArray.append(subs)
        for subcode in subcodeList:
            #if subcode:
            subcodeArray.append(subcode.strip(" "))
#print subArray #stores array of all subject names
############################################################print subcodeArray


"""
        teacher = arr[3].strip()
        #print teacher
        Tlist=teacher.splitlines()
        for teachers in Tlist:
            teacherArray.append(teachers)
        #for t in teacher:
        #teacherArray.append(teacher)

            #print teacherArray
        #print arr
       # print(sub_code, "Subject:" +sub_name, "Teacher:"+teacher)
    #print("")
# for info in time_tables:
    # print(info)
#print subArray
#print teacherArray
#print list(set(teacherArray))


#print text
"""
for table in wordDoc.tables:
    class_room = "5a"
    for row in table.rows:
        for cell in row.cells:
            time_tables.append(cell.text)
#print time_tables
TABLE = iterutils.chunked(time_tables , 16) #creates a 2D list from the available 1D list
#print TABLE

"""
for i in range(len(TABLE)):
    print(TABLE[i])
"""

# print TABLE[2]
# print len(TABLE)
slotDictionary = {'8.00-9.00':'1' , '9.00-10.00':'2' , '10.00-11.00':'3','11.00-11.30':'4', '11.30-12.30':'5' , '12.30-1.30':'6' , '1.30-2.00':'7' , '2.00-3.00' :'8' , '3.00-4.00':'9' , '4.00-5.00':'10', '8.00-11.00':'21', '11.00-2.00':'22' , '2.00-5.00':'23'}
faculty5a= {'ME':'Mahalakshmi C V','CN':'J Girija','DBMS':'Savitha S K', 'ATCI':'Sushma H R','AJAVA':'T.VijayaKumar'}
# print slotsdict.values()
# print TABLE[0]
# for i in range(1,16):
#     print(TABLE[0][i])

mon = []
mon=TABLE[1]
#print mon
for i in range(0,9):
    for j in range(0,16):
        pattern = re.search(r'\d\d\d',TABLE[i][j])
        #matches=pattern.search(TABLE[i][j])
        # if matches:
        #     print matches
        if pattern:
            print str(pattern)

# matchResult = mydog.search(s)
# >>> if matchResult:
#         print(matchResult.group(1))


# for i in range(1,9):
#     for j in range(1,16):                         #sorted...slot number and day!!!!!!!
#         time_slot = TABLE[0][j]
#         if (TABLE[i][j]!= " "):
#             print TABLE[i][j] + " is at " + TABLE[0][j]
#             print "Slot number of " +TABLE[i][j] + " is " + slotDictionary[TABLE[0][j]]
#             print "day of " +TABLE[i][j] + " is " + TABLE[i][0]               #displaying likr this for convinience sake XD
#



# def generate_mapid():
#
#
#     for i in range(0,16):
#
#
#
#     return "5a"+ mapid

#generate_mapid()
